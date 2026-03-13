"""
ai_analyzer.py — Motor de IA para análise de evidências de auditoria
Versão 4.1 — Modo Econômico integrado

Dois perfis de custo selecionáveis:
  COMPLETO  (padrão): gpt-4o, imagens detail=high, todos os docs, 3500 tokens out
  ECONÔMICO (novo)  : gpt-4o-mini, imagens detail=low, limite de páginas, 1500 tokens out
                      ≈ 10x mais barato — análise ainda sólida para decisões de auditoria
"""

import os
import re
import sys
import base64
import json
from pathlib import Path
from typing import Optional, List, Dict, Any

# ──────────────────────────────────────────────────────────────────────────────
# INTEGRAÇÃO COM BANCO DE DADOS (PARA APRENDIZADO)
# ──────────────────────────────────────────────────────────────────────────────
try:
    from backend.db import listar_aprendizados, buscar_contexto_relevante
except ImportError:
    def listar_aprendizados(*args, **kwargs): return []
    def buscar_contexto_relevante(*args, **kwargs): return ""

# ──────────────────────────────────────────────────────────────────────────────
# IMPORTAR CRITÉRIOS OFICIAIS
# ──────────────────────────────────────────────────────────────────────────────
try:
    from criterios_oficiais import CRITERIOS, REGRAS_GERAIS, get_criterio
    _CRITERIOS_DISPONIVEIS = True
except ImportError:
    _CRITERIOS_DISPONIVEIS = False
    CRITERIOS = {}
    REGRAS_GERAIS = ""
    def get_criterio(p, s): return {}

# ──────────────────────────────────────────────────────────────────────────────
# LIMITES DE TOKENS
# Perfil COMPLETO  : gpt-4o,      contexto 128k, ideal para análise profunda
# Perfil ECONÔMICO : gpt-4o-mini, contexto 128k, ~10x mais barato por token
# ──────────────────────────────────────────────────────────────────────────────
# --- PERFIL COMPLETO (padrão) ---
MAX_CHARS_TOTAL_DOCS   = 400_000   # total de texto de documentos
MAX_CHARS_PDF_PAGINA   = 3_000     # chars por página de PDF
MAX_PAGES_PDF          = 50        # páginas por PDF (todas na prática)
MAX_CHARS_DOCX         = 6_000     # chars por DOCX
MAX_CHARS_XLSX         = 4_000     # chars por Excel
MAX_IMGS_HIGH_DETAIL   = 30        # imagens em detail="high" (até 30; acima usa low)
# --- PERFIL ECONÔMICO ---
ECO_MAX_CHARS_TOTAL    = 80_000    # 5× menos texto total
ECO_MAX_CHARS_PDF      = 1_500     # 2 páginas equiv.
ECO_MAX_PAGES_PDF      = 8         # máx 8 páginas por PDF
ECO_MAX_CHARS_DOCX     = 2_500
ECO_MAX_CHARS_XLSX     = 2_000
ECO_MAX_IMGS           = 12        # no máximo 12 imagens, sempre detail="low"
ECO_MODEL              = "gpt-4o-mini"
FULL_MODEL             = "gpt-4o"


# ──────────────────────────────────────────────────────────────────────────────
# EXTRATORES DE CONTEÚDO
# ──────────────────────────────────────────────────────────────────────────────

def extract_pdf_text(file_path: str, max_chars: int = MAX_CHARS_PDF_PAGINA * MAX_PAGES_PDF) -> str:
    """Extrai texto de PDF usando pdfplumber — TODAS as páginas."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                t = page.extract_text()
                if t and t.strip():
                    text_parts.append(f"[Pág. {i+1}/{total_pages}]\n{t.strip()}")
        full = "\n".join(text_parts)
        if len(full) > max_chars:
            # Truncar preservando início e fim (mais relevantes em documentos técnicos)
            half = max_chars // 2
            full = full[:half] + f"\n\n[... {len(full)-max_chars} chars omitidos para caber no contexto ...]\n\n" + full[-half:]
        return full if full.strip() else f"[PDF sem texto extraível: {Path(file_path).name}]"
    except ImportError:
        return f"[pdfplumber não instalado — pip install pdfplumber]"
    except Exception as e:
        return f"[Erro ao extrair PDF '{Path(file_path).name}': {e}]"


def extract_docx_text(file_path: str, max_chars: int = MAX_CHARS_DOCX) -> str:
    """Extrai texto de arquivo Word (.docx) — parágrafos + tabelas."""
    try:
        from docx import Document
        doc = Document(file_path)
        parts = []
        # Parágrafos
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        # Tabelas (resumo)
        for t_idx, table in enumerate(doc.tables[:10]):
            rows_text = []
            for row in table.rows[:20]:  # máx 20 linhas por tabela
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    rows_text.append(" | ".join(cells))
            if rows_text:
                parts.append(f"[Tabela {t_idx+1}]\n" + "\n".join(rows_text))
        full = "\n".join(parts)
        return full[:max_chars] if len(full) > max_chars else full
    except ImportError:
        return f"[python-docx não instalado — pip install python-docx]"
    except Exception as e:
        return f"[Erro ao extrair DOCX '{Path(file_path).name}': {e}]"


def extract_xlsx_content(file_path: str, max_chars: int = MAX_CHARS_XLSX) -> str:
    """Extrai conteúdo real de planilha Excel (.xlsx/.xls) com openpyxl."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames[:5]:  # máx 5 abas
            ws = wb[sheet_name]
            rows_data = []
            max_row  = min(ws.max_row  or 0, 200)  # máx 200 linhas
            max_col  = min(ws.max_column or 0, 30)  # máx 30 colunas
            for row in ws.iter_rows(min_row=1, max_row=max_row,
                                    min_col=1, max_col=max_col, values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                # Ignorar linhas completamente vazias
                if any(c.strip() for c in cells):
                    rows_data.append(" | ".join(cells))
            if rows_data:
                parts.append(f"[Aba: {sheet_name}]\n" + "\n".join(rows_data))
        wb.close()
        full = "\n\n".join(parts)
        return full[:max_chars] if len(full) > max_chars else full
    except ImportError:
        # Fallback: tentar xlrd para .xls
        try:
            import xlrd
            wb = xlrd.open_workbook(file_path)
            parts = []
            for sheet in wb.sheets()[:3]:
                rows_data = []
                for ridx in range(min(sheet.nrows, 150)):
                    row = [str(sheet.cell_value(ridx, c)) for c in range(min(sheet.ncols, 20))]
                    if any(v.strip() for v in row):
                        rows_data.append(" | ".join(row))
                if rows_data:
                    parts.append(f"[Aba: {sheet.name}]\n" + "\n".join(rows_data))
            full = "\n\n".join(parts)
            return full[:max_chars] if len(full) > max_chars else full
        except Exception:
            return f"[Excel: {Path(file_path).name} — instale openpyxl para extrair conteúdo]"
    except Exception as e:
        return f"[Erro ao extrair Excel '{Path(file_path).name}': {e}]"


def image_to_base64(file_path: str) -> Optional[str]:
    """Converte imagem para base64."""
    try:
        with open(file_path, "rb") as f:
            data = f.read()
        return base64.b64encode(data).decode("utf-8")
    except Exception:
        return None


def get_image_mime(file_path: str) -> str:
    ext = Path(file_path).suffix.lower().lstrip(".")
    return {
        "jpg": "image/jpeg", "jpeg": "image/jpeg",
        "png": "image/png",  "gif": "image/gif",
        "webp": "image/webp","bmp": "image/png",
    }.get(ext, "image/jpeg")


# ──────────────────────────────────────────────────────────────────────────────
# MAPEAMENTO DE PASTAS → SUBITENS
# ──────────────────────────────────────────────────────────────────────────────

def parse_subitem_code(folder_name: str) -> Optional[tuple]:
    """
    Extrai (pratica, subitem) do nome da pasta.
    Ex: '1.3. TESTE BACKUP...' → (1, 3)
        '2.1 VERIFICAÇÃO...'   → (2, 1)
    """
    m = re.match(r'^(\d+)[\.\s]+(\d+)', folder_name.strip())
    if m:
        return (int(m.group(1)), int(m.group(2)))
    return None

# ──────────────────────────────────────────────────────────────────────────────
# PREPARAÇÃO ADAPTATIVA DE EVIDÊNCIAS
# ──────────────────────────────────────────────────────────────────────────────

def preparar_evidencias(
    evidence_files: List[str],
    max_chars_total: int = MAX_CHARS_TOTAL_DOCS,
    max_chars_pdf:   int = MAX_CHARS_PDF_PAGINA * MAX_PAGES_PDF,
    max_pages_pdf:   int = MAX_PAGES_PDF,
    max_chars_docx:  int = MAX_CHARS_DOCX,
    max_chars_xlsx:  int = MAX_CHARS_XLSX,
) -> Any:
    """
    Processa arquivos de evidência de forma adaptativa.
    Aceita parâmetros de limite para suportar modo econômico.

    Retorna:
        (evidencias_textuais: list[dict], image_files: list[Path], relatorio: str)
    """
    image_exts = {'.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp'}
    text_exts  = {'.pdf', '.docx', '.doc', '.xlsx', '.xls'}

    image_files = [Path(f) for f in evidence_files
                   if Path(f).suffix.lower() in image_exts]
    text_files  = [Path(f) for f in evidence_files
                   if Path(f).suffix.lower() in text_exts]

    total_docs = len(text_files)

    # ── Calcular budget de chars por documento ──────────────────
    if total_docs == 0:
        chars_por_doc = 0
    elif total_docs == 1:
        chars_por_doc = max_chars_total
    elif total_docs <= 5:
        chars_por_doc = max_chars_total // total_docs
    elif total_docs <= 10:
        chars_por_doc = max_chars_total // total_docs
    elif total_docs <= 20:
        chars_por_doc = max_chars_total // total_docs
    else:
        chars_por_doc = max_chars_total // total_docs

    # ── Extrair texto de TODOS os documentos ───────────────────
    evidencias_textuais = []
    erros_extracao = []

    for fp in text_files:
        ext  = fp.suffix.lower()
        nome = fp.name

        try:
            if ext == '.pdf':
                conteudo = extract_pdf_text(str(fp), max_chars=min(chars_por_doc, 100_000))
                tipo = "PDF"
            elif ext in ('.docx', '.doc'):
                conteudo = extract_docx_text(str(fp), max_chars=chars_por_doc)
                tipo = "Word"
            elif ext in ('.xlsx', '.xls'):
                conteudo = extract_xlsx_content(str(fp), max_chars=chars_por_doc)
                tipo = "Excel"
            else:
                conteudo = f"[Arquivo não suportado para extração de texto: {nome}]"
                tipo = "Outro"

            evidencias_textuais.append({
                'nome':     nome,
                'tipo':     tipo,
                'conteudo': conteudo,
            })
        except Exception as e:
            erros_extracao.append(nome)
            evidencias_textuais.append({
                'nome':     nome,
                'tipo':     'Erro',
                'conteudo': f"[Falha ao processar '{nome}': {e}]",
            })

    # ── Relatório de cobertura ──────────────────────────────────
    total_geral = len(evidence_files)
    total_chars = sum(len(ev['conteudo']) for ev in evidencias_textuais)
    relatorio_partes = [
        f"📊 COBERTURA DE ANÁLISE: {total_geral} arquivo(s) total",
        f"   📄 {total_docs} documento(s) textual(is) → {total_chars:,} chars extraídos",
        f"   🖼️ {len(image_files)} imagem(ns)",
    ]
    if erros_extracao:
        relatorio_partes.append(f"   ⚠️ {len(erros_extracao)} erro(s) de extração: {', '.join(erros_extracao)}")
    relatorio = "\n".join(relatorio_partes)

    return evidencias_textuais, image_files, relatorio


# ──────────────────────────────────────────────────────────────────────────────
# SISTEMA DE PROMPT
# ──────────────────────────────────────────────────────────────────────────────

SYSTEM_PROMPT = f"""Você é um sistema de Auditoria de TI/TA especializado em conformidade com o procedimento PO.AUT.002 Rev3.
Você opera no modo 'SQUAD DE ESPECIALISTAS' para garantir o máximo rigor técnico:

1. 👨‍💻 AUDITOR LÍDER: Mantém o foco absoluto no PO.AUT.002. Não aceita justificativas vagas. Exige provas concretas para cada nível de maturidade (0-4).
2. 🔍 ANALISTA DE EVIDÊNCIAS: Especialista em análise forense. Cruza dados de logs, nomes de arquivos, datas e metadados visuais para confirmar a autenticidade.
3. ⚖️ ESPECIALISTA EM RISCO E COMPLIANCE: Avalia se a "Prática" descrita condiz com a "Evidência" apresentada. Identifica lacunas que impedem notas superiores.

DIRETRIZES FUNDAMENTAIS:
- CRITÉRIO TEMPORAL: Evidências devem ter < 12 meses. Se ultrapassado, nota 1 ou 0 dependendo da gravidade.
- CRITÉRIO SAP: Menção a ordens de manutenção exige comprovação visual da ordem ou registro textual inequívoco.
- RIGOR TÉCNICO: Não assuma que algo existe se não estiver documentado. "Dúvida do Auditor" = Nota Inferior.
- SEMPRE responda em Português (Brasil).
- Responda EXCLUSIVAMENTE em JSON."""


def _montar_niveis_texto(niveis_oficiais: dict, niveis_planilha: dict) -> str:
    nomes = ["Não tem prática", "Iniciando", "Regular", "Bom", "Excelente"]
    linhas = []
    for k in range(5):
        txt_oficial = niveis_oficiais.get(k, "")
        txt_plan    = niveis_planilha.get(k, "")
        txt = txt_oficial or txt_plan
        if txt:
            linhas.append(f"NOTA {k} — {nomes[k]}: {txt}")
    return "\n".join(linhas)


# ──────────────────────────────────────────────────────────────────────────────
# MAPEAMENTO DE EVIDÊNCIAS
# ──────────────────────────────────────────────────────────────────────────────

def build_evidence_map(ev_folder: str) -> dict:
    """
    Escaneia a pasta de evidências e mapeia arquivos para (pratica_num, subitem_idx).
    Mesma lógica do frontend para garantir consistência.
    """
    mapa = {}
    if not ev_folder or not os.path.isdir(ev_folder):
        return mapa

    root = Path(ev_folder)
    # Extensões suportadas (conforme frontend)
    EXTS_ALL = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.pdf', '.xlsx', '.xls', '.docx', '.doc', '.mp4', '.avi', '.mov', '.mkv', '.webm'}

    try:
        # Padrão: [Prática XX] Nome / X.Y Nome
        for pasta_pratica in sorted(root.iterdir()):
            if not pasta_pratica.is_dir():
                continue
            
            # Extrair número da prática: [01] ou [1] ou 1 ou 01
            m_p = re.match(r'^\[?(\d+)[\]\-_\s\.]', pasta_pratica.name)
            if not m_p:
                continue
            p_num = int(m_p.group(1))

            for pasta_sub in sorted(pasta_pratica.iterdir()):
                if not pasta_sub.is_dir():
                    continue
                
                # Extrair número do subitem: 1.1 ou 1.01
                m_s = re.match(r'^(\d+)\.(\d+)(?:[\s\-_\.]|$)', pasta_sub.name)
                if not m_s:
                    continue
                s_num = int(m_s.group(2)) - 1 # 0-based index

                files = [
                    f for f in sorted(pasta_sub.rglob("*"))
                    if f.is_file() and f.suffix.lower() in EXTS_ALL
                ]
                if files:
                    mapa[(p_num, s_num)] = files
    except Exception as e:
        print(f"Erro ao mapear evidências: {e}")

    return mapa


def build_audit_prompt(
    pratica_num: int,
    subitem_idx: int,
    pratica_nome: str,
    subitem_nome: str,
    evidencia_descricao_planilha: str,
    niveis_planilha: dict,
    nota_self_assessment: int,
    evidencias_textuais: list,
    num_imagens: int,
    relatorio_cobertura: str = "",
    image_names: list = None,
    aprendizados: str = "",
    contexto_extra: str = "",
    regimento_geral: str = ""
) -> str:
    """Constrói o prompt de auditoria com TODOS os documentos extraídos e contexto local."""

    criterio_oficial       = get_criterio(pratica_num, subitem_idx) if _CRITERIOS_DISPONIVEIS else {}
    niveis_oficiais        = criterio_oficial.get("niveis", {})
    evidencia_desc_oficial = criterio_oficial.get("evidencias_exigidas", "")
    regras_especiais       = criterio_oficial.get("regras_especiais", "")

    evidencia_desc = evidencia_desc_oficial or evidencia_descricao_planilha
    niveis_txt     = _montar_niveis_texto(niveis_oficiais, niveis_planilha)

    # ── Formatar evidências textuais ──────────────────────────────
    if evidencias_textuais:
        ev_partes = []
        for idx, ev in enumerate(evidencias_textuais, 1):
            ev_partes.append(
                f"{'─'*60}\n"
                f"EVIDÊNCIA {idx}/{len(evidencias_textuais)}: {ev['nome']} ({ev['tipo']})\n"
                f"{'─'*60}\n"
                f"{ev['conteudo']}"
            )
        ev_texto = "\n\n".join(ev_partes)
    else:
        ev_texto = "Nenhum documento textual disponível."

    # Inventário de imagens com nomes de arquivo
    if num_imagens > 0 and image_names:
        _inv_lines = []
        for _ii, _iname in enumerate(image_names, 1):
            _inv_lines.append(f"  [{_ii:02d}] {_iname}")
        img_info = (
            f"{num_imagens} imagem(ns) enviada(s) para análise visual direta.\n"
            f"INVENTÁRIO DE IMAGENS (ordem de envio):\n"
            + "\n".join(_inv_lines)
        )
    elif num_imagens > 0:
        img_info = f"{num_imagens} imagem(ns) enviada(s) para análise visual direta."
    else:
        img_info = "Nenhuma imagem enviada."

    secao_regras = (
        f"\n⚠️ REGRAS ESPECÍFICAS DESTE SUBITEM (obrigatório considerar):\n{regras_especiais}\n"
        if regras_especiais else ""
    )

    nota_declarada = nota_self_assessment if nota_self_assessment is not None else "não informada"
    cobertura_txt  = f"\n{relatorio_cobertura}\n" if relatorio_cobertura else ""

    return f"""═══════════════════════════════════════════════════════════
OBJETO DA AUDITORIA: {subitem_nome}
REFERÊNCIA: {pratica_nome} (P{pratica_num}.{subitem_idx + 1})
═══════════════════════════════════════════════════════════
{cobertura_txt}
{contexto_extra}
{aprendizados}

CRITÉRIO DE EVIDÊNCIA (O que deve ser apresentado):
{evidencia_desc if evidencia_desc else '(consultar PO.AUT.002)'}

REGRAS ESPECÍFICAS:
{regras_especiais if regras_especiais else 'Nenhuma regra especial disparada.'}

NÍVEIS DE MATURIDADE (Critério de Pontuação):
{niveis_txt}

CONTEXTO DO DISPOSITIVO (Self Assessment):
A área declarou nota {nota_declarada}/4. Sua missão é validar se os arquivos abaixo sustentam essa nota.

═══════════════════════════════════════════════════════════
ANÁLISE DE EVIDÊNCIAS DOCUMENTAIS ({len(evidencias_textuais)} docs)
═══════════════════════════════════════════════════════════
{ev_texto}

═══════════════════════════════════════════════════════════
ANÁLISE DE EVIDÊNCIAS VISUAIS
═══════════════════════════════════════════════════════════
{img_info}

═══════════════════════════════════════════════════════════
MISSÃO DO SQUAD:
1. Analise cada documento/imagem fundo. Identifique datas e tags de ativos.
2. Aplique os critérios de maturidade com "Pé no Chão". Não conceda nota por "promessa" de melhoria.
3. Determine se o status é:
   - "permanece": Evidência COMPLETA e ATUAL que justifica a nota {nota_declarada}.
   - "insuficiente": Evidência existe mas é parcial, desatualizada ou incompleta para a nota {nota_declarada}.
   - "inexistente": Não há prova material válida. Nota deve ser 0.

RESPOSTA OBRIGATÓRIA EM JSON:
{{
  "decisao": "permanece" | "insuficiente" | "inexistente",
  "nota_sugerida": <int 0-4>,
  "confianca": "alta" | "media" | "baixa",
  "pontos_atendidos": ["Lista de fatos positivos encontrados"],
  "pontos_faltantes": ["O que falta para atingir a nota desejada ou a nota 4"],
  "descricao_nc": "Relato técnico da Não Conformidade ou GAP encontrado",
  "comentarios": "Sugestões de melhoria para o próximo ciclo",
  "analise_detalhada": "Justificativa técnica profunda mencionando nomes de arquivos e dados específicos extraídos.",
  "inventario_imagens": [
    {{ "arquivo": "...", "equipamento_identificado": "...", "valida": true, "observacao": "..." }}
  ]
}}"""


# ──────────────────────────────────────────────────────────────────────────────
# ANALISADOR PRINCIPAL
# ──────────────────────────────────────────────────────────────────────────────

class AuditAIAnalyzer:
    def __init__(self, api_key: str, model: str = FULL_MODEL, economico: bool = False, 
                 provider: str = "openai", base_url: Optional[str] = None):
        self.api_key   = api_key
        self.economico = economico
        self.provider  = str(provider or "openai").lower()
        self.base_url  = base_url
        # Modo econômico força gpt-4o-mini se for OpenAI
        if self.provider == "openai":
            self.model = ECO_MODEL if economico else (model or FULL_MODEL)
        elif self.provider == "ollama":
            self.model = model or "auditoria-ta"
        else:
            self.model = model

    def _get_client(self):
        if self.provider == "openai":
            from openai import OpenAI
            return OpenAI(api_key=self.api_key, base_url=self.base_url)
        elif self.provider == "ollama":
            # Ollama costuma usar a interface da OpenAI ou uma biblioteca própria.
            # Vamos assumir compatibilidade com API OpenAI (via /v1) se base_url for provido.
            from openai import OpenAI
            return OpenAI(api_key="ollama", base_url=self.base_url or "http://localhost:11434/v1")
        elif self.provider == "gemini":
            # Para Gemini via API da Google
            import google.generativeai as genai
            genai.configure(api_key=self.api_key)
            return genai
        elif self.provider == "anthropic":
            from anthropic import Anthropic
            return Anthropic(api_key=self.api_key, base_url=self.base_url)
        return None

    def analyze_subitem(
        self,
        pratica_num: int,
        subitem_idx: int,
        pratica_nome: str,
        subitem_nome: str,
        evidencia_descricao: str,
        niveis_planilha: dict,
        nota_self_assessment: int,
        evidence_files: list
    ) -> dict:
        """
        Analisa um subitem processando TODOS os arquivos de evidência.

        Estratégia adaptativa:
          - Documentos textuais: extrai texto de TODOS com budget dinâmico por doc
          - Imagens: detail="high" se ≤8, "low" se >8 (cobre até 100 imagens)
          - Excel: extração real do conteúdo das células
          - Relatório de cobertura incluído no prompt
        """
        # ── Preparar todas as evidências ──────────────────────────
        if self.economico:
            evidencias_textuais, image_paths, relatorio_cobertura = preparar_evidencias(
                evidence_files,
                max_chars_total = ECO_MAX_CHARS_TOTAL,
                max_chars_pdf   = ECO_MAX_CHARS_PDF,
                max_pages_pdf   = ECO_MAX_PAGES_PDF,
                max_chars_docx  = ECO_MAX_CHARS_DOCX,
                max_chars_xlsx  = ECO_MAX_CHARS_XLSX,
            )
            image_paths = image_paths[:ECO_MAX_IMGS]
        else:
            evidencias_textuais, image_paths, relatorio_cobertura = preparar_evidencias(evidence_files)

        # ── Injeção de Aprendizado (Few-Shot) ──────────────────────
        aprendizados = listar_aprendizados(pratica_num, subitem_idx)
        exemplos_aprendizado = []
        if aprendizados:
            exemplos_aprendizado = []
            for apr in aprendizados[:3]: # Usar os 3 mais recentes
                ex = apr.get("exemplo", "")
                desc = apr.get("descricao", "")
                exemplos_aprendizado.append(f"CASO ANTERIOR: {desc}\nSOLUÇÃO/APRENDIZADO: {ex}")
            contexto_aprendizado = "💡 LIÇÕES APRENDIDAS EM CASOS ANTERIORES SEMELHANTES:\n" + \
                                   "\n".join(exemplos_aprendizado)

        # ── Construir prompt ───────────────────────────────────────
        # 4. Gerar Prompt
        prompt = build_audit_prompt(
            pratica_num                  = pratica_num,
            subitem_idx                  = subitem_idx,
            pratica_nome                 = pratica_nome,
            subitem_nome                 = subitem_nome,
            evidencia_descricao_planilha = evidencia_descricao,
            niveis_planilha              = niveis_planilha,
            nota_self_assessment         = nota_self_assessment,
            evidencias_textuais          = evidencias_textuais,
            num_imagens                  = len(image_paths),
            relatorio_cobertura          = relatorio_cobertura,
            image_names                  = [p.name for p in image_paths],
        )

        # Injetar exemplos de aprendizado no prompt
        if exemplos_aprendizado:
            prompt = "💡 LIÇÕES APRENDIDAS EM CASOS ANTERIORES SEMELHANTES:\n" + \
                     "\n".join(exemplos_aprendizado) + \
                     "\n\n" + prompt

        # ── Montar mensagem multimodal com imagens ─────────────────
        # Modo econômico: sempre detail="low", máx tokens de saída menor
        if self.economico:
            img_detail  = "low"
            max_tok_out = 1500
        else:
            img_detail  = "high" if len(image_paths) <= MAX_IMGS_HIGH_DETAIL else "low"
            max_tok_out = 3500
        # IMPORTANTE: cada imagem é precedida por rótulo de texto com seu nome
        content    = [{"type": "text", "text": prompt}]
        total_imgs = len(image_paths)

        imgs_adicionadas = 0
        imgs_erro = []
        for img_idx, img_path in enumerate(image_paths, 1):
            b64 = image_to_base64(str(img_path))
            if b64:
                mime = get_image_mime(str(img_path))
                # Rótulo de texto ANTES da imagem: IA sabe o nome do arquivo
                content.append({
                    "type": "text",
                    "text": f"📸 IMAGEM {img_idx}/{total_imgs} — {img_path.name}"
                })
                content.append({
                    "type": "image_url",
                    "image_url": {
                        "url":    f"data:{mime};base64,{b64}",
                        "detail": img_detail,
                    }
                })
                imgs_adicionadas += 1
            else:
                imgs_erro.append(img_path.name)

        # ── Chamada à API ──────────────────────────────────────────
        result_raw = ""
        client = self._get_client()
        
        if self.provider in ("openai", "ollama"):
            response = client.chat.completions.create(
                model    = self.model,
                messages = [
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user",   "content": content},
                ],
                max_tokens       = max_tok_out,
                temperature      = 0.1,
                response_format  = {"type": "json_object"} if self.provider == "openai" else None,
            )
            result_raw = response.choices[0].message.content
        
        elif self.provider == "gemini":
            # Gemini support
            model = client.GenerativeModel(self.model)
            # Convert multimodal content to Gemini format
            gemini_content = []
            for item in content:
                if item["type"] == "text":
                    gemini_content.append(item["text"])
                elif item["type"] == "image_url":
                    # Gemini expects raw bytes for images in some SDK versions, 
                    # but easiest is to pass the b64 data or parts.
                    # This is a simplification.
                    pass 
            
            # For now, let's focus on text-only if complex multimodal is too different, 
            # or use the proper Gemini multimodal structure.
            # Simplified Gemini call:
            response = model.generate_content([SYSTEM_PROMPT] + gemini_content)
            result_raw = response.text
        elif self.provider == "anthropic":
            # Claude 3.5 Sonnet / Haiku
            # Anthropic usa um formato de mensagens um pouco diferente para multimodal
            anthropic_messages = []
            msg_content = [{"type": "text", "text": prompt}]
            
            for img_path in image_paths[:MAX_IMGS_HIGH_DETAIL]:
                # Claude exige base64 e media_type
                with open(img_path, "rb") as f:
                    data = base64.b64encode(f.read()).decode("utf-8")
                    media_type = "image/jpeg" if img_path.suffix.lower() == ".jpg" else "image/png"
                    msg_content.append({
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": data
                        }
                    })
            
            anthropic_messages.append({"role": "user", "content": msg_content})
            
            response = client.messages.create(
                model=self.model or "claude-3-5-sonnet-20240620",
                max_tokens=max_tok_out,
                system=SYSTEM_PROMPT,
                messages=anthropic_messages
            )
            result_raw = response.content[0].text

        # Limpar markdown do JSON se necessário
        print(f"DEBUG: Resposta bruta da IA ({len(result_raw)} chars)")
        
        try:
            if "```json" in result_raw:
                result_raw = result_raw.split("```json")[-1].split("```")[0].strip()
            elif "```" in result_raw:
                result_raw = result_raw.split("```")[-1].split("```")[0].strip()
            
            result = json.loads(result_raw)
        except Exception as e:
            print(f"ERROR: Falha ao parsear JSON da IA: {e}")
            print(f"CONTEUDO BRUTO QUE FALHOU: {result_raw[:500]}...")
            # Fallback seguro para não estourar 500
            result = {
                "decisao": "insuficiente",
                "nota_sugerida": 0,
                "confianca": "baixa",
                "analise_detalhada": f"ERRO DE PARSE JSON: {str(e)}\n\nResposta bruta: {result_raw[:1000]}"
            }

        # ── Validar e normalizar resultado ─────────────────────────
        decisoes_validas = {'permanece', 'insuficiente', 'inexistente'}
        if result.get('decisao') not in decisoes_validas:
            result['decisao'] = 'insuficiente'

        nota_sug = result.get('nota_sugerida')
        if not isinstance(nota_sug, (int, float)) or nota_sug < 0 or nota_sug > 4:
            nota_map = {
                'permanece':    nota_self_assessment,
                'insuficiente': max(0, (nota_self_assessment or 0) - 1),
                'inexistente':  0,
            }
            result['nota_sugerida'] = nota_map.get(result['decisao'], nota_self_assessment)
        else:
            result['nota_sugerida'] = int(nota_sug)

        # ── Aplicar regras hard ────────────────────────────────────
        try:
            result = self._aplicar_regras_hard(
                pratica_num, subitem_idx, result, nota_self_assessment, evidence_files
            )
        except Exception as e:
            print(f"ERROR: Falha ao aplicar regras hard: {e}")

        # ── Metadados de cobertura ─────────────────────────────────
        result['nota_self_assessment'] = nota_self_assessment
        result['arquivos_analisados']  = len(evidence_files)
        result['imagens_analisadas']   = imgs_adicionadas
        result['docs_analisados']      = len(evidencias_textuais)
        result['imgs_detalhe']         = img_detail
        result['criterios_source']     = 'oficial' if _CRITERIOS_DISPONIVEIS else 'planilha'
        result['cobertura_relatorio']  = relatorio_cobertura

        # ── Processar inventário de imagens ───────────────────────
        # Serializa o inventario_imagens para texto e ANEXA à analise_detalhada
        # (assim é salvo no banco sem precisar de nova coluna)
        inv_imgs = result.get('inventario_imagens', [])
        if inv_imgs and isinstance(inv_imgs, list):
            try:
                inv_linhas = ["\n\n━━━ INVENTÁRIO DE IMAGENS ANALISADAS ━━━"]
                for item in inv_imgs:
                    arq   = item.get('arquivo', '?')
                    equip = item.get('equipamento_identificado', '?')
                    valida= '✅' if item.get('valida', True) else '⚠️'
                    obs   = item.get('observacao', '')
                    inv_linhas.append(f"{valida} {arq} → {equip}" + (f" | {obs}" if obs else ""))
                inv_txt = "\n".join(inv_linhas)
                # Anexar ao analise_detalhada existente
                analise_atual = result.get('analise_detalhada', '') or ''
                result['analise_detalhada'] = analise_atual + inv_txt
            except Exception:
                pass  # Não quebrar se formato vier inesperado

        return result

    def _aplicar_regras_hard(self, pratica_num: int, subitem_idx: int,
                              result: dict, nota_sa: int, evidence_files: list) -> dict:
        """
        Aplica regras hard que resultam em nota 0 obrigatória,
        independente do parecer da IA.
        """
        if not _CRITERIOS_DISPONIVEIS:
            return result

        criterio       = get_criterio(pratica_num, subitem_idx)
        regras_hard    = criterio.get("regras_hard", [])
        image_exts     = {'.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp'}
        text_exts      = {'.pdf', '.docx', '.doc', '.xlsx', '.xls'}
        n_imgs         = sum(1 for f in evidence_files if Path(f).suffix.lower() in image_exts)
        n_docs         = sum(1 for f in evidence_files if Path(f).suffix.lower() in text_exts)
        total_evidencias = len(evidence_files)

        for regra in regras_hard:
            tipo = regra.get("tipo")

            if tipo == "sem_evidencia_fisica" and total_evidencias == 0:
                result['decisao']      = 'inexistente'
                result['nota_sugerida'] = 0
                result['confianca']    = 'alta'
                # Assegurar que pontos_faltantes seja uma lista
                if not isinstance(result.get('pontos_faltantes'), list):
                    result['pontos_faltantes'] = []
                result['pontos_faltantes'].append("[REGRA HARD] Sem nenhuma evidência física anexada")
                break

            if tipo == "exige_doc_especifico":
                nome_exigido = regra.get("nome_arquivo", "").lower()
                encontrado = any(
                    nome_exigido in Path(f).name.lower()
                    for f in evidence_files
                )
                if not encontrado:
                    result['decisao']      = 'inexistente'
                    result['nota_sugerida'] = 0
                    result['confianca']    = 'alta'
                    # Assegurar que pontos_faltantes seja uma lista
                    if not isinstance(result.get('pontos_faltantes'), list):
                        result['pontos_faltantes'] = []
                    result['pontos_faltantes'].append(
                        f"[REGRA HARD] Documento obrigatório não encontrado: {regra.get('nome_arquivo')}"
                    )
                    break

        return result

    def revisar_com_contexto(
        self,
        pratica_num: int,
        subitem_idx: int,
        pratica_nome: str,
        subitem_nome: str,
        evidencia_descricao: str,
        niveis_planilha: dict,
        nota_self_assessment: int,
        evidence_files: list,
        historico_chat: list,
        nova_mensagem: str,
        exemplos_aprendizado: list = None,
    ) -> dict:
        """
        Revisão colaborativa humano-IA: mantém contexto da conversa e
        incorpora nova observação do auditor humano.

        Retorna dict com:
          - resposta: str
          - mudou_decisao: bool
          - nova_decisao: str | None
          - nova_nota: int | None
          - nova_confianca: str | None
        """
        from openai import OpenAI

        # ── Construir system prompt para revisão ──────────────────
        try:
            from criterios_oficiais import get_criterio as _gc, REGRAS_GERAIS as _RG
            criterio          = _gc(pratica_num, subitem_idx)
            niveis_oficiais   = criterio.get("niveis", {})
            evidencia_exigida = criterio.get("evidencias_exigidas", "")
            regras_especiais  = criterio.get("regras_especiais", "")
            regras_gerais_txt = _RG
        except Exception:
            niveis_oficiais   = {}
            evidencia_exigida = ""
            regras_especiais  = ""
            regras_gerais_txt = ""

        nomes_nv   = ["Não tem prática", "Iniciando", "Regular", "Bom", "Excelente"]
        niveis_txt = ""
        for k in range(5):
            txt = niveis_oficiais.get(k, niveis_planilha.get(k, ""))
            if txt:
                niveis_txt += f"\nNOTA {k} — {nomes_nv[k]}: {txt}"

        system_revisao = f"""Você é um auditor especialista em automação industrial.
Você fez uma análise inicial e agora está em revisão colaborativa com o auditor humano.

CONTEXTO DO SUBITEM:
Prática: {pratica_nome}
Subitem: {subitem_nome} (P{pratica_num}.{subitem_idx+1})
Nota Self Assessment declarada: {nota_self_assessment}/4
Evidências exigidas: {evidencia_exigida or '(ver critérios)'}
{f'Regras especiais: {regras_especiais}' if regras_especiais else ''}

CRITÉRIOS OFICIAIS (PO.AUT.002):
{niveis_txt}

REGRAS GERAIS:
{regras_gerais_txt}

INSTRUÇÃO PARA REVISÃO:
- O auditor pode apresentar observações adicionais, apontar divergências ou confirmar seu parecer.
- Se as novas informações mudarem sua avaliação, indique EXPLICITAMENTE: REVISÃO_DECISAO: <decisao> | REVISÃO_NOTA: <nota> | REVISÃO_CONFIANÇA: <confiança>
- Se mantiver sua posição, explique tecnicamente o motivo com base nos critérios oficiais.
- Sempre considere que o auditor de campo tem visibilidade adicional das evidências.
- Seja objetivo, técnico e colaborativo.
- Responda em português do Brasil."""

        # ── Injetar exemplos de aprendizado (few-shot) ────────────
        if exemplos_aprendizado:
            exemplos_txt = "\n\nCASOS ANTERIORES SEMELHANTES (aprendizado acumulado):\n"
            for i, ex in enumerate(exemplos_aprendizado, 1):
                exemplos_txt += f"""
Caso {i} ({ex.get('unidade','?')} — {ex.get('ciclo','?')}):
  SA={ex.get('nota_sa','?')} | IA inicial={ex.get('nota_ia_inicial','?')} ({ex.get('decisao_ia_inicial','?')})
  Observação do auditor: {ex.get('observacao_auditor','?')}
  Consenso final: {ex.get('decisao_consenso','?')} | Nota={ex.get('nota_consenso','?')}
  Justificativa: {ex.get('justificativa','') or '—'}
"""
            system_revisao = system_revisao + exemplos_txt

        # ── Montar histórico ──────────────────────────────────────
        mensagens = [{"role": "system", "content": system_revisao}]

        for msg in historico_chat:
            if msg["role"] in ("user", "assistant"):
                mensagens.append({
                    "role":    msg["role"],
                    "content": msg.get("content") or msg.get("conteudo", ""),
                })

        mensagens.append({"role": "user", "content": nova_mensagem})

        # ── Chamada à API ──────────────────────────────────────────
        client   = OpenAI(api_key=self.api_key)
        response = client.chat.completions.create(
            model       = self.model,
            messages    = mensagens,
            max_tokens  = 1200,
            temperature = 0.2,
        )

        resposta_texto = response.choices[0].message.content

        # ── Detectar revisão de decisão ────────────────────────────
        import re as _re
        mudou       = False
        nova_decisao = None
        nova_nota    = None
        nova_conf    = None

        m_dec = _re.search(
            r"REVIS[AÃ]O_DECIS[AÃ]O:\s*(permanece|insuficiente|inexistente)",
            resposta_texto, _re.IGNORECASE)
        m_dec_raw = _re.search(
            r"REVIS[AÃ]O_DECIS[AÃ]O:\s*([^|\n]+)",
            resposta_texto, _re.IGNORECASE)
        m_nota = _re.search(r"REVIS[AÃ]O_NOTA:\s*(\d)", resposta_texto, _re.IGNORECASE)
        m_conf = _re.search(r"REVIS[AÃ]O_CONFIAN[CÇ]A:\s*(alta|media|m[eé]dia|baixa)",
                            resposta_texto, _re.IGNORECASE)

        if m_dec:
            nova_decisao = m_dec.group(1).lower()
            mudou = True
        elif m_dec_raw:
            _raw  = m_dec_raw.group(1).strip().lower()
            _mapa = {
                "ajustar": "insuficiente", "ajuste":   "insuficiente",
                "reduz":   "insuficiente", "reduzir":  "insuficiente",
                "insuf":   "insuficiente", "incompleto":"insuficiente",
                "parcial": "insuficiente", "inadequado":"insuficiente",
                "inexist": "inexistente",  "zero":      "inexistente",
                "sem evid":"inexistente",  "mant":      "permanece",
                "confirm": "permanece",    "ok":        "permanece",
            }
            for chave, valor in _mapa.items():
                if chave in _raw:
                    nova_decisao = valor
                    mudou = True
                    break

        if m_nota:
            nova_nota = int(m_nota.group(1))
            mudou = True
        if m_conf:
            _conf_raw = m_conf.group(1).lower()
            nova_conf = "media" if "dia" in _conf_raw else _conf_raw

        # Nota prevalece sobre decisão em conflito
        if nova_nota is not None and nova_decisao is None:
            if nova_nota == 0:
                nova_decisao = "inexistente"; mudou = True
            elif nova_nota < nota_self_assessment:
                nova_decisao = "insuficiente"; mudou = True
            else:
                nova_decisao = "permanece"

        if nova_nota is not None and nova_decisao is not None:
            if nova_nota == 0 and nova_decisao != "inexistente":
                nova_decisao = "inexistente"; mudou = True
            elif nova_nota > 0 and nova_nota < nota_self_assessment and nova_decisao == "permanece":
                nova_decisao = "insuficiente"; mudou = True

        return {
            "resposta":       resposta_texto,
            "mudou_decisao":  mudou,
            "nova_decisao":   nova_decisao,
            "nova_nota":      nova_nota,
            "nova_confianca": nova_conf,
        }

    def analyze_batch(self, subitems_data: list, evidence_map: dict,
                      progress_callback=None) -> list:
        """
        Analisa múltiplos subitens em sequência.
        subitems_data: lista de dicts com campos do subitem.
        """
        results = []
        total   = len(subitems_data)

        for i, sub in enumerate(subitems_data):
            key   = (sub['pratica_num'], sub['subitem_idx'])
            files = evidence_map.get(key, [])

            try:
                result = self.analyze_subitem(
                    pratica_num          = sub['pratica_num'],
                    subitem_idx          = sub['subitem_idx'],
                    pratica_nome         = sub['pratica_nome'],
                    subitem_nome         = sub['subitem_nome'],
                    evidencia_descricao  = sub.get('evidencia_descricao', ''),
                    niveis_planilha      = sub.get('niveis', {}),
                    nota_self_assessment = sub.get('nota_sa', 0) or 0,
                    evidence_files       = files,
                )
                result['pratica_num']  = sub['pratica_num']
                result['subitem_idx']  = sub['subitem_idx']
                result['subitem_nome'] = sub['subitem_nome']
                result['status']       = 'ok'
            except Exception as e:
                result = {
                    'pratica_num':       sub['pratica_num'],
                    'subitem_idx':       sub['subitem_idx'],
                    'subitem_nome':      sub['subitem_nome'],
                    'status':            'erro',
                    'erro':              str(e),
                    'decisao':           'pendente',
                    'nota_sugerida':     sub.get('nota_sa'),
                    'confianca':         'baixa',
                    'pontos_atendidos':  [],
                    'pontos_faltantes':  [],
                    'descricao_nc':      '',
                    'comentarios':       f'[Erro na análise de IA: {e}]',
                    'analise_detalhada': '',
                }

            results.append(result)
            if progress_callback:
                progress_callback(i + 1, total, sub['subitem_nome'])

        return results


# ──────────────────────────────────────────────────────────────────────────────
# VALIDAÇÃO RÁPIDA DE CHAVE API
# ──────────────────────────────────────────────────────────────────────────────

def validar_chave_openai(api_key: str) -> dict:
    """
    Testa a chave da OpenAI com uma chamada mínima (1 token).
    Retorna dict com campos:
      valida       (bool)   – chave aceita pela API
      erro         (str)    – mensagem de erro amigável, ou ""
      tipo_erro    (str)    – "sem_credito" | "chave_invalida" | "sem_acesso" | "outro" | ""
      modelo_ok    (bool)   – True se gpt-4o-mini respondeu OK
    """
    if not api_key or not api_key.strip().startswith("sk-"):
        return {"valida": False, "erro": "Chave inválida: deve começar com 'sk-'", "tipo_erro": "chave_invalida", "modelo_ok": False}
    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key.strip())
        client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": "ping"}],
            max_tokens=1,
            temperature=0,
        )
        return {"valida": True, "erro": "", "tipo_erro": "", "modelo_ok": True}
    except Exception as e:
        err = str(e)
        if "insufficient_quota" in err or "429" in err:
            return {
                "valida": False,
                "erro": (
                    "❌ Cota esgotada ou orçamento do projeto zerado.\n\n"
                    "Causas mais comuns:\n"
                    "• O projeto OpenAI tem orçamento = $0 → acesse platform.openai.com → "
                    "seu projeto → Limits → ajuste o Monthly budget\n"
                    "• Saldo da conta zerado → Billing → Add to credit balance"
                ),
                "tipo_erro": "sem_credito",
                "modelo_ok": False,
            }
        elif "invalid_api_key" in err or "Incorrect API key" in err or "401" in err:
            return {
                "valida": False,
                "erro": "❌ Chave inválida. Copie a chave novamente em platform.openai.com → API Keys.",
                "tipo_erro": "chave_invalida",
                "modelo_ok": False,
            }
        elif "model_not_found" in err or "does not have access" in err:
            return {
                "valida": False,
                "erro": "❌ Sua chave não tem acesso ao modelo solicitado.",
                "tipo_erro": "sem_acesso",
                "modelo_ok": False,
            }
        else:
            return {
                "valida": False,
                "erro": f"❌ Erro ao testar chave: {err}",
                "tipo_erro": "outro",
                "modelo_ok": False,
            }
